# pip install langgraph langchain openai

from typing import TypedDict
from langgraph.graph import StateGraph, START, END
from langchain_openai import ChatOpenAI
from docx import Document
# -----------------------------
# 1. Define State
# -----------------------------
class BlogState(TypedDict):
    topic: str
    outline: str
    blog: str


# -----------------------------
# 2. Initialize LLM
# -----------------------------
llm = ChatOpenAI(
    model="gpt-4o-mini",
    api_key="sk-proj-W-SFptWvRAZq5R2p2139V_tA42f27OBYI8PSftZA4CYYADEAOaAhc   your API KEY here",
    temperature=0.7
)


# -----------------------------
# 3. Node 1: Create Outline
# -----------------------------
def create_outline(state: BlogState):
    topic = state["topic"]

    prompt = f"Create a detailed blog outline for the topic: {topic}"

    response = llm.invoke(prompt)

    return {
        "outline": response.content
    }


# -----------------------------
# 4. Node 2: Create Blog
# -----------------------------
def create_blog(state: BlogState):
    topic = state["topic"]
    outline = state["outline"]

    prompt = f"""
    Write a full blog based on the topic and outline.

    Topic: {topic}

    Outline:
    {outline}
    """

    response = llm.invoke(prompt)

    return {
        "blog": response.content
    }


# -----------------------------
# 5. Build Graph
# -----------------------------
graph = StateGraph(BlogState)

# Add nodes
graph.add_node("create_outline", create_outline)
graph.add_node("create_blog", create_blog)

# Set entry point
graph.add_edge(START, "create_outline")

# Add edges (sequential flow)
graph.add_edge("create_outline", "create_blog")
graph.add_edge("create_blog", END)

# Compile graph
app = graph.compile()


# -----------------------------
# 5. Save to Word Function
# -----------------------------
def save_to_word(state: BlogState, filename="blog_output.docx"):
    doc = Document()

    # Title
    doc.add_heading(state["topic"], 0)

    # Outline Section
    doc.add_heading("Outline", level=1)
    doc.add_paragraph(state["outline"])

    # Blog Section
    doc.add_heading("Blog", level=1)
    doc.add_paragraph(state["blog"])

    doc.save(filename)
    print(f"\n✅ Saved to {filename}")



# -----------------------------
# 6. Run Graph
# -----------------------------
if __name__ == "__main__":
    input_data = {
        "topic": "Benefits of Artificial Intelligence in Healthcare"
    }

    result = app.invoke(input_data)

    print("\n--- OUTLINE ---\n")
    print(result["outline"])

    print("\n--- BLOG ---\n")
    print(result["blog"])

    # Save to Word
    save_to_word(result)
